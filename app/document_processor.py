"""
Document Processor - Extract text from various file formats and chunk it.
"""

import os
import re
import hashlib
from io import BytesIO
from typing import List, Tuple
import uuid

from app.models import DocumentChunk


class DocumentProcessor:
    """
    Handles document text extraction and chunking.
    
    Supported formats:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Plain Text (.txt)
    """
    
    def __init__(self):
        self._init_extractors()
    
    def _init_extractors(self):
        """Initialize text extractors lazily."""
        pass
    
    def extract_text(self, content: bytes, extension: str, filename: str) -> str:
        """
        Extract text from document content.
        
        Args:
            content: Raw file bytes
            extension: File extension (e.g., '.pdf')
            filename: Original filename for error reporting
            
        Returns:
            Extracted text as string
        """
        extension = extension.lower()
        
        if extension == ".pdf":
            return self._extract_pdf(content, filename)
        elif extension == ".docx":
            return self._extract_docx(content, filename)
        elif extension == ".txt":
            return self._extract_txt(content, filename)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _extract_pdf(self, content: bytes, filename: str) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required for PDF processing. Install with: pip install pdfplumber")
        
        text_parts = []
        
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            text_parts.append(f"[Page {i+1}]\n{page_text}")
                        
                        # Also extract tables if present
                        tables = page.extract_tables()
                        for j, table in enumerate(tables):
                            if table:
                                table_text = self._table_to_text(table)
                                text_parts.append(f"[Table {j+1} on Page {i+1}]\n{table_text}")
                    except Exception as e:
                        print(f"Warning: Error extracting page {i+1} from {filename}: {e}")
                        continue
        except Exception as e:
            raise ValueError(f"Failed to parse PDF '{filename}': {str(e)}")
        
        if not text_parts:
            raise ValueError(f"No text could be extracted from PDF '{filename}'. The file may be scanned/image-based.")
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, content: bytes, filename: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = Document(BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX '{filename}': {str(e)}")
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            
            if table_rows:
                table_text = self._table_to_text(table_rows)
                text_parts.append(f"[Table {i+1}]\n{table_text}")
        
        if not text_parts:
            raise ValueError(f"No text could be extracted from DOCX '{filename}'.")
        
        return "\n\n".join(text_parts)
    
    def _extract_txt(self, content: bytes, filename: str) -> str:
        """Extract text from plain text file."""
        # Try common encodings
        encodings = ['utf-8', 'utf-16', 'cp1251', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                if text.strip():
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        raise ValueError(f"Could not decode text file '{filename}' with any supported encoding.")
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert a table (list of rows) to text format."""
        if not table:
            return ""
        
        lines = []
        for row in table:
            if row:
                # Filter out None values and convert to strings
                cells = [str(cell) if cell else "" for cell in row]
                lines.append(" | ".join(cells))
        
        return "\n".join(lines)
    
    def chunk_text(
        self,
        text: str,
        source_file: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks for embedding.
        
        Args:
            text: Full document text
            source_file: Source filename
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            
        Returns:
            List of DocumentChunk objects
        """
        if not text.strip():
            return []
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Split into chunks using sentence boundaries when possible
        chunks = self._smart_chunk(text, chunk_size, chunk_overlap)
        
        # Create DocumentChunk objects
        doc_chunks = []
        base_id = hashlib.md5(source_file.encode()).hexdigest()[:8]
        
        current_pos = 0
        for i, chunk_text in enumerate(chunks):
            # Find actual position in cleaned text
            start_char = text.find(chunk_text[:50], current_pos)
            if start_char == -1:
                start_char = current_pos
            end_char = start_char + len(chunk_text)
            current_pos = max(current_pos, end_char - chunk_overlap)
            
            chunk = DocumentChunk(
                chunk_id=f"{base_id}-{i}",
                text=chunk_text,
                source_file=source_file,
                chunk_index=i,
                start_char=start_char,
                end_char=end_char,
                metadata={
                    "word_count": len(chunk_text.split()),
                    "char_count": len(chunk_text),
                }
            )
            doc_chunks.append(chunk)
        
        return doc_chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Replace multiple whitespace with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _smart_chunk(
        self,
        text: str,
        chunk_size: int,
        chunk_overlap: int,
    ) -> List[str]:
        """
        Split text into chunks, trying to respect sentence boundaries.
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + chunk_size
            
            if end >= len(text):
                # Last chunk
                chunks.append(text[start:])
                break
            
            # Try to find a sentence boundary near the end
            # Look for period, question mark, or exclamation followed by space
            best_break = end
            
            # Search backward for sentence boundary
            search_start = max(start + chunk_size // 2, start)
            search_text = text[search_start:end]
            
            # Find last sentence boundary in the search range
            sentence_endings = ['. ', '? ', '! ', '.\n', '?\n', '!\n']
            last_boundary = -1
            
            for ending in sentence_endings:
                pos = search_text.rfind(ending)
                if pos > last_boundary:
                    last_boundary = pos
            
            if last_boundary > 0:
                best_break = search_start + last_boundary + 2  # Include the ending
            else:
                # Try to break at paragraph
                para_break = search_text.rfind('\n\n')
                if para_break > 0:
                    best_break = search_start + para_break + 2
                else:
                    # Try to break at any newline
                    newline = search_text.rfind('\n')
                    if newline > 0:
                        best_break = search_start + newline + 1
                    else:
                        # Break at space
                        space = search_text.rfind(' ')
                        if space > 0:
                            best_break = search_start + space + 1
            
            # Extract chunk
            chunk = text[start:best_break].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = best_break - chunk_overlap
            if start < 0:
                start = 0
        
        return chunks
    
    def detect_document_type(self, text: str) -> dict:
        """
        Analyze text to detect contract type and key sections.
        
        Returns metadata about the document structure.
        """
        text_lower = text.lower()
        
        # Detect document type
        doc_type = "unknown"
        if any(term in text_lower for term in ["purchase agreement", "договор поставки", "supply agreement"]):
            doc_type = "supply_agreement"
        elif any(term in text_lower for term in ["service agreement", "договор оказания услуг"]):
            doc_type = "service_agreement"
        elif any(term in text_lower for term in ["amendment", "дополнительное соглашение", "изменение"]):
            doc_type = "amendment"
        elif any(term in text_lower for term in ["nda", "confidentiality", "конфиденциальность"]):
            doc_type = "nda"
        elif any(term in text_lower for term in ["contract", "договор", "agreement", "соглашение"]):
            doc_type = "general_contract"
        
        # Detect sections
        sections_found = []
        section_patterns = [
            (r"article\s+\d+|статья\s+\d+|раздел\s+\d+", "articles"),
            (r"price|pricing|цена|стоимость", "pricing"),
            (r"payment|оплата|платеж", "payment"),
            (r"delivery|поставка|доставка", "delivery"),
            (r"termination|расторжение|прекращение", "termination"),
            (r"warranty|гарантия|гарантийн", "warranty"),
            (r"penalty|штраф|санкц", "penalties"),
            (r"amendment|изменен|дополнен", "amendments"),
            (r"sla|service level|уровень сервиса", "sla"),
        ]
        
        for pattern, section_name in section_patterns:
            if re.search(pattern, text_lower):
                sections_found.append(section_name)
        
        return {
            "document_type": doc_type,
            "sections_found": sections_found,
            "language": self._detect_language(text),
            "word_count": len(text.split()),
        }
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on character frequency."""
        # Count Cyrillic vs Latin characters
        cyrillic_count = len(re.findall(r'[а-яА-ЯёЁ]', text))
        latin_count = len(re.findall(r'[a-zA-Z]', text))
        
        if cyrillic_count > latin_count:
            return "ru"
        elif latin_count > 0:
            return "en"
        else:
            return "unknown"
